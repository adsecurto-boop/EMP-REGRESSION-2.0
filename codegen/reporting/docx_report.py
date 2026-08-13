"""DOCX report generator using python-docx.
"""

from pathlib import Path
from typing import Optional
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from codegen.reporting.models import ExecutionSummary, TestResult


def set_cell_background(cell, fill_color: str):
    """Set shading/background color of a table cell (hex without #)."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set internal cell padding in dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


class DOCXReportGenerator:
    """Generates professional DOCX QA regression reports."""

    @staticmethod
    def generate(summary: ExecutionSummary, output_path: Path) -> Path:
        """Generate test_report.docx for execution summary."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # 1. Title Section
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("EMPMonitor Regression Test Report")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(30, 41, 59)

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run("Automated Playwright QA Execution Report • Version 0.1.3")
        sub_run.font.name = "Arial"
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph()  # spacing

        # 2. Executive Summary
        h2 = doc.add_heading("1. Executive Summary", level=1)
        h2.style.font.color.rgb = RGBColor(30, 41, 59)

        exec_p = doc.add_paragraph()
        exec_p.add_run("This report documents the automated Playwright regression test suite execution for ").font.size = Pt(10)
        bold_app = exec_p.add_run("EmpMonitor")
        bold_app.bold = True
        bold_app.font.size = Pt(10)
        exec_p.add_run(". The results reflect all executed test cases, test data, expected versus actual outcomes, and failure evidence.").font.size = Pt(10)

        # Stats Table
        stats_table = doc.add_table(rows=2, cols=5)
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        stats_headers = ["Total Tests", "Passed", "Failed", "Skipped", "Pass Rate"]
        stats_values = [
            str(summary.total),
            str(summary.passed),
            str(summary.failed),
            str(summary.skipped),
            f"{summary.pass_percentage:.1f}%"
        ]

        hdr_cells = stats_table.rows[0].cells
        for i, h_text in enumerate(stats_headers):
            hdr_cells[i].text = h_text
            set_cell_background(hdr_cells[i], "1E293B")
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)

        val_cells = stats_table.rows[1].cells
        for i, val in enumerate(stats_values):
            val_cells[i].text = val
            set_cell_background(val_cells[i], "F8FAFC")
            p = val_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(12)
                if i == 1:
                    r.font.color.rgb = RGBColor(21, 128, 61)
                elif i == 2:
                    r.font.color.rgb = RGBColor(185, 28, 28)

        doc.add_paragraph()

        # 3. Environment Details
        h3 = doc.add_heading("2. Environment & Execution Details", level=1)
        h3.style.font.color.rgb = RGBColor(30, 41, 59)

        env_table = doc.add_table(rows=4, cols=2)
        env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        env_data = [
            ("Application / Platform", "EmpMonitor Web Dashboard"),
            ("Test Environment & Base URL", f"{summary.environment} ({summary.base_url})"),
            ("Browser & Execution Engine", f"Playwright {summary.browser}"),
            ("Execution Timestamp & Verdict", f"{summary.timestamp} — VERDICT: {summary.overall_status}")
        ]

        for idx, (label, val) in enumerate(env_data):
            row_cells = env_table.rows[idx].cells
            row_cells[0].text = label
            row_cells[1].text = val
            set_cell_background(row_cells[0], "F1F5F9")
            set_cell_background(row_cells[1], "FFFFFF")
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)

        doc.add_paragraph()

        # 4. Complete Test Results Table
        h4 = doc.add_heading("3. Complete Test Case Results", level=1)
        h4.style.font.color.rgb = RGBColor(30, 41, 59)

        # Columns: TC ID, Module, Test Case, Test Data, Expected Result, Actual Result, Status, Duration
        res_table = doc.add_table(rows=1, cols=8)
        res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["TC ID", "Module", "Test Case", "Test Data", "Expected Result", "Actual Result", "Status", "Duration"]

        hdr_row = res_table.rows[0]
        for idx, text in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.text = text
            set_cell_background(cell, "1E293B")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(8.5)

        for test in summary.tests:
            row = res_table.add_row()
            cells = row.cells
            cells[0].text = str(test.test_id)
            cells[1].text = str(test.module)
            cells[2].text = str(test.title)
            cells[3].text = str(test.test_data) if test.test_data else "-"
            cells[4].text = str(test.expected)
            cells[5].text = str(test.actual)
            cells[6].text = str(test.status)
            cells[7].text = f"{test.duration_seconds:.2f}s"

            # Style text size and background
            status_bg = "DCFCE7" if test.status == "PASS" else ("FEE2E2" if test.status == "FAIL" else "FEF3C7")
            set_cell_background(cells[6], status_bg)

            for i, cell in enumerate(cells):
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(8)
                    if i == 6:
                        r.font.bold = True
                        if test.status == "PASS":
                            r.font.color.rgb = RGBColor(21, 128, 61)
                        elif test.status == "FAIL":
                            r.font.color.rgb = RGBColor(185, 28, 28)

        doc.add_paragraph()

        # 5. Failure Evidence Section
        failed_tests = [t for t in summary.tests if t.status == "FAIL"]
        if failed_tests:
            h_fail = doc.add_heading("4. Failed Test Cases & Failure Evidence", level=1)
            h_fail.style.font.color.rgb = RGBColor(185, 28, 28)

            for ft in failed_tests:
                f_p = doc.add_paragraph()
                r_title = f_p.add_run(f"[{ft.test_id}] {ft.title} ({ft.module})")
                r_title.bold = True
                r_title.font.size = Pt(11)

                p_detail = doc.add_paragraph()
                p_detail.add_run(f"• Expected: {ft.expected}\n").font.size = Pt(9.5)
                p_detail.add_run(f"• Actual: {ft.actual}\n").font.size = Pt(9.5)
                p_detail.add_run(f"• Error: {ft.exception_message or ft.failure_reason}\n").font.size = Pt(9)

                # Embed Screenshot
                if ft.screenshot_path and Path(ft.screenshot_path).exists():
                    try:
                        doc.add_paragraph("Screenshot Evidence:").runs[0].font.bold = True
                        doc.add_picture(str(ft.screenshot_path), width=Inches(5.5))
                    except Exception as e:
                        doc.add_paragraph(f"(Could not attach screenshot image: {e})")

                doc.add_paragraph()

        # 6. Overall Verdict
        h_v = doc.add_heading("5. Overall Test Verdict", level=1)
        h_v.style.font.color.rgb = RGBColor(30, 41, 59)

        v_p = doc.add_paragraph()
        v_p.add_run("FINAL REGRESSION SUITE STATUS: ").font.size = Pt(11)
        status_run = v_p.add_run(summary.overall_status)
        status_run.bold = True
        status_run.font.size = Pt(12)
        if summary.overall_status == "PASSED":
            status_run.font.color.rgb = RGBColor(21, 128, 61)
        else:
            status_run.font.color.rgb = RGBColor(185, 28, 28)

        doc.save(str(output_path))
        return output_path
